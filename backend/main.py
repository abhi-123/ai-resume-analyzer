
from fastapi.middleware.cors import CORSMiddleware

from fastapi import APIRouter, UploadFile, File, Form,FastAPI
from PyPDF2 import PdfReader
from docx import Document
from dotenv import load_dotenv
from openai import OpenAI
import json

import os
import io
app = FastAPI()

# ✅ CORS (React/JS connect)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # 🔥 sab allow (dev ke liye)
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
# 🔑 Load env variables
load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
router = APIRouter()

@app.post("/analyze")
async def analyze(file: UploadFile = File(None), text: str = Form(None)):
    textToBeSend = ''
    if file:
        filename = file.filename
        ext = filename.split(".")[-1]
        content = await file.read()
        if ext == 'pdf': 
            reader = PdfReader(io.BytesIO(content))
            for page in reader.pages:
                textToBeSend += page.extract_text() or ""
        else:
            
            doc = Document(io.BytesIO(content))
            textToBeSend = "\n".join([para.text for para in doc.paragraphs])
    else:
     return {"error": "No input provided"}
    data = summarize_resume(textToBeSend,text)
    return data
       
def summarize_resume(text,job_description=''):
    try:
     # 🔥 safety limit (LLM token control)
        #text = text[:6000]
        system_prompt_llm = """You are an expert resume analyzer and ATS optimization specialist.

Your job is to evaluate resumes with high precision and return structured JSON output only.

If a job description is provided, you MUST strictly evaluate the resume in the context of that job role, focusing on relevance, matching skills, and missing requirements.
"""

        if job_description:
            system_prompt_llm += f"""

        Job Description:
        {job_description}
        """


        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": system_prompt_llm
                },
                {
                    "role": "user",
                    "content": f"""
        Analyze the following resume and return STRICTLY valid JSON in this format:

        {{
        "score": number (0-100),
        "summary": ["point1", "point2", "point3"],
        "strengths": ["point1", "point2"],
        "weaknesses": ["point1", "point2"],
        "suggestions": ["point1", "point2"]
        }}

        Evaluation Criteria:
        - Relevance to job role (if job description provided)
        - Clarity and structure
        - Presence of key sections: projects, work experience, skills, education
        - Use of impact-driven language and measurable achievements

        Job Description Rules (IMPORTANT):
        - If job description is provided:
        - Prioritize relevance to the role
        - Penalize missing required skills
        - Highlight mismatch clearly
        - Tailor suggestions specifically for this role

        Low Quality Resume Rules:
        - If resume is incoherent, irrelevant, or lacks meaningful professional content:
        - Set "score" to 0
        - Return ONLY:
            - "summary"
            - "suggestions"
        - Set:
            - "strengths": []
            - "weaknesses": []

        Output Rules:
        - Return ONLY JSON (no text, no explanation)
        - Keep all points:
        - Short
        - Crisp
        - Non-repetitive
        - Actionable
        - Do NOT hallucinate fake experience or skills

        Resume:
        {text}
        """
                }
                    ],
                    response_format={"type": "json_object"}
                )

        result = json.loads(response.choices[0].message.content)
        print(result)
        return {
            "success": True,
            "data": result
        }

    except Exception as e:
        return {
            "success": False,
            "error": str(e)
        }

@app.post("/rewrite")
async def rewrite(file: UploadFile = File(None), text: str = Form(None) , weaknesses: str = Form(None)):
    weaknesses_list = []
    if weaknesses:
     weaknesses_list = json.loads(weaknesses)
     weakness_text = "\n".join([f"- {w}" for w in weaknesses_list])
    textToBeSend = ''
    if file:
        filename = file.filename
        ext = filename.split(".")[-1]
        content = await file.read()
        if ext == 'pdf': 
            reader = PdfReader(io.BytesIO(content))
            for page in reader.pages:
                textToBeSend += page.extract_text() or ""
        else:
            
            doc = Document(io.BytesIO(content))
            textToBeSend = "\n".join([para.text for para in doc.paragraphs])
    elif text:
       textToBeSend = text
    else:
     return {"error": "No input provided"}
    data = rewrite_resume(textToBeSend,weakness_text,text)
    return data

def rewrite_resume(text,weakness_text,job_description=''):
    print(weakness_text)
    try:
     # 🔥 safety limit (LLM token control)
        text = text[:6000]
        system_prompt_llm = """You are an expert resume writer, ATS optimization specialist, and career coach.

    Your task is to rewrite resumes to make them highly professional, impactful, and tailored for ATS systems.

    If a job description is provided, you MUST strictly optimize the resume for that role by aligning skills, experience, and keywords accordingly.
    """

        if job_description:
            system_prompt_llm += f"""

        Job Description:
        {job_description}
        """


        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": system_prompt_llm
                },
                {
                    "role": "user",
                    "content": f"""
        Rewrite the following resume to make it highly professional, ATS-optimized, and impactful.

        Weaknesses to fix:
        {weakness_text}

        ------------------------------------

        Instructions:

        1. Improve Content Quality:
        - Use strong action verbs (e.g., Developed, Led, Optimized)
        - Add measurable impact (numbers, %, scale) where possible
        - Make each point concise and results-driven
        - Avoid repetition and fluff

        2. ATS Optimization:
        - Include relevant keywords naturally
        - Ensure clarity in skills and experience
        - Improve readability and structure

        3. Job Description Alignment (CRITICAL):
        - If job description is provided:
        - Align experience and skills with the role
        - Prioritize relevant technologies and responsibilities
        - Address missing skills indirectly where possible
        - Do NOT invent fake experience

        4. Constraints:
        - Do NOT hallucinate or fabricate experience
        - Do NOT add technologies not implied by the resume
        - Keep all points short, crisp, and professional

        5. Summary:
        - Write 2–3 strong, impactful bullet points
        - Focus on experience, specialization, and value

        6. "changed" Field:
        - Clearly describe what was improved:
        - e.g., "Added quantification, improved action verbs, aligned with JD, enhanced clarity"

        ------------------------------------

        Return STRICTLY valid JSON (no extra text):

        {{
        "summary": ["point1", "point2", "point3"],
        "experience": ["point1", "point2", "point3"],
        "projects": ["point1", "point2", "point3"],
        "skills": ["skill1", "skill2", "skill3"],
        "changed": "string explaining improvements"
        }}

        ------------------------------------

        Resume:
        {text}
        """
                }
                ],
                response_format={"type": "json_object"}
            )

        result = json.loads(response.choices[0].message.content)
        print(result)
        return {
            "success": True,
            "data": result
        }

    except Exception as e:
        return {
            "success": False,
            "error": str(e)
        }

